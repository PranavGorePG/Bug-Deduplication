import re
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("python-docx not installed. Run 'pip install python-docx'")
    exit(1)

def add_markdown_paragraph(paragraph, text):
    # minimal markdown parser for **bold** and *italic* and `code`
    # This is a simple implementation.
    
    # Split by bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Split by code
            subparts = re.split(r'(`.*?`)', part)
            for subpart in subparts:
                if subpart.startswith('`') and subpart.endswith('`'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.font.name = 'Courier New'
                    # subtle background color not easily possible without xml manipulation in basic python-docx
                    # so we just use font
                else:
                    paragraph.add_run(subpart)

def create_doc(md_path, doc_path):
    doc = Document()
    
    # Add Title
    # doc.add_heading('System Architecture', 0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        line_stripped = line.strip()
        
        if line_stripped.startswith('```'):
            if in_code_block:
                # End of code block
                if code_block_content:
                    # Add a table with one cell to simulate a code block box
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)
                    p = cell.paragraphs[0]
                    p.style = 'No Spacing'
                    run = p.add_run('\n'.join(code_block_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                
                in_code_block = False
                code_block_content = []
            else:
                # Start of code block
                in_code_block = True
            continue
            
        if in_code_block:
            code_block_content.append(line.rstrip()) # keep indentation but remove newline
            continue
            
        # Strip trailing newlines for processing, but keep indentation if it's a list item?
        # Markdown lists usually don't depend on indentation as much as bullet characters
        # But sublists do. For simplicity, we'll just strip.
        line = line.strip()
        
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_markdown_paragraph(p, line[2:])
        elif re.match(r'^\d+\.', line):
             p = doc.add_paragraph(style='List Number')
             text = re.sub(r'^\d+\.\s*', '', line)
             add_markdown_paragraph(p, text)
        else:
            p = doc.add_paragraph()
            add_markdown_paragraph(p, line)
            
    doc.save(doc_path)
    print(f"Document saved to {doc_path}")

md_file = r"C:\Users\hp\.gemini\antigravity\brain\c2b2a870-2930-455b-b2d2-953c8d14dd02\system_architecture.md"
doc_file = r"C:\Users\hp\.gemini\antigravity\brain\c2b2a870-2930-455b-b2d2-953c8d14dd02\system_architecture.docx"

if __name__ == "__main__":
    create_doc(md_file, doc_file)
